"""
양식 파일 없이 추가 요청사항 기반으로 마스터 DB 열을 그대로 추출하는 스크립트

동작:
  - AI가 notes를 읽고 필요한 열(db_index)과 필터 조건을 결정
  - 마스터 DB에서 해당 열을 그대로 추출 → Excel로 저장
"""
import json, os, sys, datetime, re
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

sys.path.insert(0, os.path.dirname(__file__))
from common import load_master_db, filter_by_year, to_yyyymmdd, get_year, safe_int
from ai_client import create_json_response


# ── DB 전체 컬럼 목록 (AI 참조용) ──────────────────────────────────
DB_COLUMNS_SUMMARY = """
마스터 DB 컬럼 목록 (db_index → 헤더명):
0=연번, 1=기술이전계약일, 2=기술보유기관명, 3=기술도입업체명,
4=기관유형(코드), 5=업종유형, 6=국내외(1=국내,2=국외), 7=국가명,
8=국내지역구분(지역코드), 9=사업자등록번호, 10=대표주소, 11=대표전화,
12=대표자성명, 13=홈페이지, 14=우편주소, 15=기술이전담당자명,
16=담당부서, 17=직급, 18=핸드폰, 19=팩스, 20=이메일,
21=산단담당자명, 22=산단담당부서, 23=산단전화번호, 24=산단이메일,
25=종업원수(상시), 26=연매출액(천원), 27=기술명, 28=기술명(27.기술명과 동일계열),
29=주발명자, 30=교직원번호, 31=주발명자소속, 32=전공,
33=공동발명자, 34=공동발명자교직원번호, 35=공동발명자소속, 36=공동발명자전공,
37=기술유형(코드), 38=지식재산권번호, 39=출원등록상태(1=출원,2=등록),
40=포함기술수, 41=특허비용, 42=기술분야(6T), 43=기술분류,
44=거래유형(코드), 45=계약기간, 46=계약시작일, 47=계약종료일,
48=제한사항, 49=기술료수취유형코드, 50=선급기술료(원), 51=경상기술료조건,
52=정액기술료(원), 53=총기술료(원), 54=계약상태(0=유지,1=해지,2=종료),
55=계약해지사유, 56=연구과제명, 57=부처명, 58=지원기관, 59=지원사업,
60=총연구비(천원), 61=총연구기간, 62=협약일, 63=대사업명, 64=중사업명,
65=지원기관과제번호, 66=ERP과제번호, 67=연구책임자, 68=공동연구자,
69=참여기업명, 70=정부출연금, 71=기업부담금, 72=기타,
73=입금일, 74=입금통장명의, 75=기술료수취유형,
76=현금입금액(원), 77=주식(원), 78=현물(원), 79=기타(원), 80=입금종류,
81=선급기술료분배액, 82=경상기술료분배액, 83=정액기술료분배액,
84=분배일, 85=제반비용_특허비용, 86=제반비용_중개수수료,
87=전문기관분배액, 88=발명자분배액, 89=발명자지정기관분배액,
90=산학협력단분배액, 91=기술이전사업화경비, 92=성과활용기여자보상금,
93=연구개발재투자, 94=기타특허지분액, 95=해당사업단명,
96=수납상황, 97=학진승인여부, 98=NTB등록여부, 99=기술가치평가여부,
100=계약변경일, 101=계약해지일, 102=납부기한, 103=담당자, 104=담당자(연구원)
"""

# ── 헤더명 매핑 (출력 시 사용) ─────────────────────────────────────
HEADER_MAP = {
    0: "연번", 1: "기술이전계약일", 2: "기술보유기관명", 3: "기술도입업체명",
    4: "기관유형", 5: "업종유형", 6: "국내/국외", 7: "국가명",
    8: "국내지역구분", 9: "사업자등록번호", 10: "대표주소", 11: "대표전화",
    12: "대표자성명", 13: "홈페이지", 14: "우편주소", 15: "기술이전담당자명",
    16: "담당부서", 17: "직급", 18: "핸드폰", 19: "팩스", 20: "이메일",
    21: "산단담당자명", 22: "산단담당부서", 23: "산단전화번호", 24: "산단이메일",
    25: "종업원수(상시)", 26: "연매출액(천원)", 27: "기술명", 28: "기술명",
    29: "주발명자", 30: "교직원번호", 31: "주발명자소속", 32: "전공",
    33: "공동발명자", 34: "공동발명자교직원번호", 35: "공동발명자소속", 36: "공동발명자전공",
    37: "기술유형", 38: "지식재산권번호", 39: "출원/등록상태",
    40: "포함기술수", 41: "특허비용", 42: "기술분야(6T)", 43: "기술분류",
    44: "거래유형", 45: "계약기간", 46: "계약시작일", 47: "계약종료일",
    48: "제한사항", 49: "기술료수취유형", 50: "선급기술료(원)", 51: "경상기술료조건",
    52: "정액기술료(원)", 53: "총기술료(원)", 54: "계약상태",
    55: "계약해지사유", 56: "연구과제명", 57: "부처명", 58: "지원기관", 59: "지원사업",
    60: "총연구비(천원)", 61: "총연구기간", 62: "협약일", 63: "대사업명", 64: "중사업명",
    65: "지원기관과제번호", 66: "ERP과제번호", 67: "연구책임자", 68: "공동연구자",
    69: "참여기업명", 70: "정부출연금", 71: "기업부담금", 72: "기타",
    73: "입금일", 74: "입금통장명의", 75: "기술료수취유형",
    76: "현금입금액(원)", 77: "주식(원)", 78: "현물(원)", 79: "기타(원)", 80: "입금종류",
    81: "선급기술료분배액", 82: "경상기술료분배액", 83: "정액기술료분배액",
    84: "분배일", 85: "제반비용(특허비용)", 86: "제반비용(중개수수료)",
    87: "전문기관분배액", 88: "발명자분배액", 89: "발명자지정기관분배액",
    90: "산학협력단분배액", 91: "기술이전사업화경비", 92: "성과활용기여자보상금",
    93: "연구개발재투자", 94: "기타(특허지분액)", 95: "해당사업단명",
    96: "수납상황", 97: "학진승인여부", 98: "NTB등록여부", 99: "기술가치평가여부",
    100: "계약변경일", 101: "계약해지일", 102: "납부기한", 103: "담당자", 104: "담당자(연구원)",
}

# 기본 추출 열 (요청사항이 없거나 불명확할 때)
DEFAULT_COLUMNS = [0, 1, 3, 28, 29, 37, 44, 50, 52, 73, 76]
SPREADSHEET_EXTENSIONS = (".xlsx", ".xlsm", ".xltx", ".xls")
DOCUMENT_EXTENSIONS = (".pdf", ".docx")


def is_spreadsheet_file(path: str) -> bool:
    import os
    return os.path.splitext(str(path).lower())[1] in SPREADSHEET_EXTENSIONS


def extract_pdf_text(path: str, max_pages: int = 20, max_chars: int = 12000) -> str:
    try:
        from pypdf import PdfReader
    except Exception as error:
        raise RuntimeError("PDF 요청파일을 읽으려면 pypdf 패키지가 필요합니다.") from error

    reader = PdfReader(path)
    parts = []
    for page_index, page in enumerate(reader.pages[:max_pages], 1):
        text = page.extract_text() or ""
        if text.strip():
            parts.append(f"[페이지 {page_index}]\n{text.strip()}")
        if sum(len(part) for part in parts) >= max_chars:
            break
    return "\n\n".join(parts)[:max_chars]


def extract_docx_text(path: str, max_chars: int = 12000) -> str:
    try:
        from docx import Document
    except Exception as error:
        raise RuntimeError("Word 요청파일을 읽으려면 python-docx 패키지가 필요합니다.") from error

    document = Document(path)
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table_index, table in enumerate(document.tables, 1):
        parts.append(f"[표 {table_index}]")
        for row in table.rows[:60]:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(values):
                parts.append(" | ".join(values))
    return "\n".join(parts)[:max_chars]


def build_request_files_context(request_file_paths=None, request_file_names=None) -> str:
    request_file_paths = request_file_paths or []
    request_file_names = request_file_names or []
    parts = []
    for idx, path in enumerate(request_file_paths, 1):
        name = request_file_names[idx - 1] if idx - 1 < len(request_file_names) else os.path.basename(path)
        ext = os.path.splitext(str(path).lower())[1]
        parts.append(f"\n## 요청파일 {idx}: {name}")
        if is_spreadsheet_file(path):
            from openpyxl import load_workbook
            wb = load_workbook(path, data_only=True, read_only=True)
            for ws in wb.worksheets[:5]:
                parts.append(f"### 시트: {ws.title}")
                for r in range(1, min(ws.max_row or 0, 12) + 1):
                    values = []
                    for c in range(1, min(ws.max_column or 0, 12) + 1):
                        value = ws.cell(r, c).value
                        if value not in (None, ""):
                            values.append(f"[열{c}]{str(value).strip()}")
                    if values:
                        parts.append(f"행{r}: " + ", ".join(values))
            wb.close()
        elif ext == ".pdf":
            parts.append("### 문서 유형: PDF")
            parts.append(extract_pdf_text(path))
        elif ext == ".docx":
            parts.append("### 문서 유형: Word")
            parts.append(extract_docx_text(path))
    return "\n".join(parts)

REQUEST_COLUMN_ALIASES = [
    ("계약일", 1),
    ("기술이전계약일", 1),
    ("업체명", 3),
    ("기관(업체)명", 3),
    ("사업자등록번호", 9),
    ("기술명", 28),
    ("발명자명", 29),
    ("주발명자", 29),
    ("성명", 29),
    ("교직원번호", 30),
    ("교번", 30),
    ("학과", 31),
    ("소속", 31),
    ("기술유형", 37),
    ("지식재산권", 38),
    ("거래유형", 44),
    ("정액기술료", 52),
    ("계약금액", 52),
    ("계약입금일", 73),
    ("입금일", 73),
    ("현금입금액", 76),
    ("입금액", 76),
    ("경상기술료", 82),
    ("정액기술료입금", 83),
]


def infer_requested_columns(notes: str) -> list:
    """AI가 열 목록을 놓쳤을 때 요청 문장에 적힌 항목 순서대로 컬럼을 추정."""
    text = notes or ""
    found = []
    for keyword, db_index in REQUEST_COLUMN_ALIASES:
        if keyword in text and db_index not in found:
            found.append(db_index)
    return found


def describe_columns(columns: list) -> list:
    """화면 검증용 컬럼 설명."""
    return [
        {"db_index": db_idx, "header": HEADER_MAP.get(db_idx, f"col_{db_idx}")}
        for db_idx in columns
    ]


def ask_openai(notes: str, year: int, api_key: str, request_files_context: str = "") -> dict:
    """
    추가 요청사항을 분석해서 추출할 열과 필터 조건을 반환.
    반환 형식:
    {
      "columns": [0, 1, 3, ...],
      "filter_mode": "contract"|"payment"|"both"|"none",
      "filter_year": 2025,
      "date_range": {"start": "YYYYMMDD", "end": "YYYYMMDD"} 또는 null,
      "researcher_filter": ["이름1", "이름2", ...] 또는 null,
      "reference_filters": [{"source_file_index": 1, "column": 2, "db_indices": [30, 34]}],
      "db_filters": [{"db_index": 76, "operator": "gte", "value": 100000000, "value_type": "number"}],
      "notes": "분석 설명"
    }
    """
    prompt = f"""당신은 부산대학교 산학협력단의 기술이전 담당자입니다.
아래 추가 요청사항을 분석해서 마스터 DB에서 어떤 열을 추출할지 결정해주세요.

## 기준 연도: {year}년

## 마스터 DB 컬럼 목록:
{DB_COLUMNS_SUMMARY}

## 사용자 추가 요청사항:
{notes if notes.strip() else "(없음 - 기본 주요 항목을 추출)"}

## 요청파일 참고 내용:
{request_files_context if request_files_context.strip() else "(없음)"}

## 응답 형식 (JSON만 출력, 다른 텍스트 없이):
{{
  "columns": [정수 db_index 목록],
  "filter_mode": "contract" 또는 "payment" 또는 "both" 또는 "none",
  "filter_year": 연도_정수_또는_null,
  "date_range": {{"start": "YYYYMMDD", "end": "YYYYMMDD"}} 또는 null,
  "researcher_filter": ["이름1", "이름2"] 또는 null,
  "reference_filters": [
    {{
      "source_file_index": 1,
      "sheet": "평가명단",
      "column": 2,
      "data_start_row": 2,
      "db_indices": [30, 34],
      "match_type": "contains",
      "label": "교직원번호"
    }}
  ],
  "db_filters": [
    {{
      "db_index": 76,
      "operator": "gte",
      "value": 100000000,
      "value_type": "number",
      "label": "현금입금액 1억원 이상"
    }}
  ],
  "notes": "선택 이유 한 줄"
}}

## 판단 기준:
- 요청사항에 특정 열이 언급되면 해당 열 포함
- 연도 필터: 계약일 기준이면 "contract", 입금일 기준이면 "payment", 둘 다면 "both", 연도 무관이면 "none"
- filter_year: 요청에 다른 연도가 있으면 그 연도, 없으면 null (기준 연도 {year} 사용)
- date_range: "X월Y일~X월Y일", "YYYY년MM월DD일부터" 등 구체적 날짜 범위가 있으면 YYYYMMDD 형식으로 변환. 없으면 null
- researcher_filter: 특정 연구자/발명자 이름 목록이 있으면 배열로 추출. 없으면 null
- reference_filters: 요청파일이 명단/대상 파일이고 "동일 대상", "명단 기준" 같은 요청이 있으면 사용. 없으면 빈 배열 [].
  - source_file_index: 요청파일 번호(1부터 시작)
  - column: 요청파일에서 대상값이 있는 열 번호(1-based)
  - data_start_row: 대상값 시작 행
  - db_indices: 마스터 DB에서 비교할 컬럼. 교직원번호 [30,34], 연구자명/발명자명 [29,33], 사업자등록번호 [9], 업체명 [3], 기술명 [28]
  - match_type: exact 또는 contains
- db_filters: 자연어 요청에 마스터 DB 조건이 있으면 구조화. 없으면 빈 배열 [].
  - operator: eq, neq, in, not_in, contains, not_contains, gt, gte, lt, lte, between, date_between, year_eq, is_empty, not_empty
  - value_type: text, number, date, auto
  - 예: "현금입금액 1억원 이상" → db_index 76 / gte / 100000000 / number
  - 예: "부처명이 교육부" → db_index 57 / contains / "교육부" / text
  - 예: "기술자문 제외" → db_index 44 / not_in / ["9", "기술자문"] / text
  - 예: "계약일 20260101~20260630" → db_index 1 / date_between / ["20260101", "20260630"] / date
- 추출 항목이 명시되지 않으면 columns는 빈 배열 [] 로 반환 (전체 열 추출로 처리됨)
- 추출 항목이 명시된 경우에는 사용자가 요청한 항목만 columns에 넣으세요. 기본 컬럼을 추가하지 마세요.
- 열 순서는 사용자가 적은 요청 항목 순서를 최대한 유지하세요.
- columns에는 중복 없이 정수만"""

    return create_json_response(
        api_key=api_key,
        prompt=prompt,
        instructions=(
            "너는 기술이전 마스터 DB 추출 조건을 JSON으로 설계하는 도우미다. "
            "반드시 유효한 JSON 객체만 출력한다."
        ),
        max_output_tokens=1800,
    )


def filter_by_date_range(rows, start_yyyymmdd: str, end_yyyymmdd: str,
                         contract_col=1, payment_col=73, mode="both") -> list:
    """날짜 범위(YYYYMMDD 문자열)로 행 필터링."""
    def to_str(v) -> str:
        return to_yyyymmdd(v) if v else ""

    result = []
    for row in rows:
        c_str = to_str(row[contract_col] if contract_col < len(row) else None)
        p_str = to_str(row[payment_col]  if payment_col  < len(row) else None)

        def in_range(s):
            return s and start_yyyymmdd <= s <= end_yyyymmdd

        if mode == "contract":
            if in_range(c_str): result.append(row)
        elif mode == "payment":
            if in_range(p_str): result.append(row)
        else:  # both
            if in_range(c_str) or in_range(p_str): result.append(row)
    return result


def filter_by_researchers(rows, names: list,
                           inventor_col=29, co_inventor_col=33) -> list:
    """주발명자 또는 공동발명자가 names 목록에 포함된 행만 반환."""
    name_set = set(n.strip() for n in names if n.strip())

    def matches(row):
        inv    = str(row[inventor_col]    or "").strip() if inventor_col    < len(row) else ""
        co_inv = str(row[co_inventor_col] or "").strip() if co_inventor_col < len(row) else ""
        # 공동발명자는 여러 명이 구분자로 이어질 수 있음
        if inv in name_set:
            return True
        for part in co_inv.replace(",", " ").replace("/", " ").replace(";", " ").split():
            if part.strip() in name_set:
                return True
        return False

    return [row for row in rows if matches(row)]


def _norm_match_value(value):
    if value is None:
        return ""
    s = str(value).strip()
    if re.fullmatch(r"\d+\.0", s):
        s = s[:-2]
    return re.sub(r"[\s\-_/]", "", s)


def _db_filter_values(db_filter):
    if "values" in db_filter and db_filter.get("values") is not None:
        raw = db_filter.get("values")
    elif "value" in db_filter:
        raw = [db_filter.get("value")]
    else:
        raw = []
    if not isinstance(raw, list):
        raw = [raw]
    return [v for v in raw if v is not None and str(v).strip() != ""]


def _compare_as_number(value):
    try:
        if value is None or str(value).strip() == "":
            return None
        return float(str(value).replace(",", "").strip())
    except Exception:
        return None


def _compare_as_date_text(value):
    text = to_yyyymmdd(value)
    if not text:
        return ""
    return re.sub(r"\D", "", str(text))[:8]


def _cell_matches_filter(cell_value, db_filter):
    operator = str(db_filter.get("operator") or "eq").lower()
    values = _db_filter_values(db_filter)
    value_type = str(db_filter.get("value_type") or db_filter.get("type") or "auto").lower()

    if value_type == "number" or operator in ("gt", "gte", "lt", "lte"):
        cell_num = _compare_as_number(cell_value)
        nums = [_compare_as_number(v) for v in values]
        nums = [n for n in nums if n is not None]
        if cell_num is None:
            return False
        if operator == "gt":
            return bool(nums) and cell_num > nums[0]
        if operator == "gte":
            return bool(nums) and cell_num >= nums[0]
        if operator == "lt":
            return bool(nums) and cell_num < nums[0]
        if operator == "lte":
            return bool(nums) and cell_num <= nums[0]
        if operator == "between":
            return len(nums) >= 2 and min(nums[0], nums[1]) <= cell_num <= max(nums[0], nums[1])
        if operator in ("eq", "in"):
            return cell_num in nums
        if operator in ("neq", "not_in"):
            return cell_num not in nums

    if value_type in ("date", "yyyymmdd") or operator in ("date_between", "year_eq"):
        cell_date = _compare_as_date_text(cell_value)
        date_values = [_compare_as_date_text(v) for v in values]
        date_values = [v for v in date_values if v]
        if not cell_date:
            return False
        if operator in ("between", "date_between"):
            return len(date_values) >= 2 and min(date_values[0], date_values[1]) <= cell_date <= max(date_values[0], date_values[1])
        if operator == "year_eq":
            years = [str(v)[:4] for v in values if str(v).strip()]
            return cell_date[:4] in years
        if operator in ("eq", "in"):
            return cell_date in date_values
        if operator in ("neq", "not_in"):
            return cell_date not in date_values

    cell_text = str(cell_value).strip() if cell_value is not None else ""
    cell_norm = _norm_match_value(cell_value)
    value_texts = [str(v).strip() for v in values]
    value_norms = [_norm_match_value(v) for v in values]

    if operator in ("eq", "equals", "in"):
        return any(cell_text == v or cell_norm == n for v, n in zip(value_texts, value_norms))
    if operator in ("neq", "not_equals", "not_in"):
        return not any(cell_text == v or cell_norm == n for v, n in zip(value_texts, value_norms))
    if operator == "contains":
        return any(v and (v in cell_text or n in cell_norm) for v, n in zip(value_texts, value_norms))
    if operator == "not_contains":
        return not any(v and (v in cell_text or n in cell_norm) for v, n in zip(value_texts, value_norms))
    if operator in ("is_empty", "empty"):
        return cell_text == ""
    if operator in ("not_empty", "is_not_empty"):
        return cell_text != ""
    return False


def apply_db_filters(rows, db_filters):
    if not db_filters:
        return rows
    filtered = rows
    for db_filter in db_filters:
        try:
            db_index = int(db_filter.get("db_index"))
        except Exception:
            print(f"  DB 필터 db_index 오류: {db_filter}")
            continue
        before = len(filtered)
        label = db_filter.get("label") or f"DB[{db_index}]"
        filtered = [
            row for row in filtered
            if db_index < len(row) and _cell_matches_filter(row[db_index], db_filter)
        ]
        print(f"  DB 조건 필터 '{label}' → {before}행 → {len(filtered)}행")
    return filtered


def _extract_reference_values(reference_path, reference_filter):
    from openpyxl import load_workbook
    sheet_name = reference_filter.get("sheet")
    col = int(reference_filter.get("column") or 0)
    start_row = int(reference_filter.get("data_start_row") or 2)
    if col < 1 or not is_spreadsheet_file(reference_path):
        return []

    wb = load_workbook(reference_path, data_only=True, read_only=True)
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
    values = []
    seen = set()
    for r in range(start_row, (ws.max_row or 0) + 1):
        value = _norm_match_value(ws.cell(r, col).value)
        if value and value not in seen:
            values.append(value)
            seen.add(value)
    wb.close()
    return values


def infer_reference_filters(request_file_paths, notes=""):
    if not request_file_paths:
        return []
    wants_reference_filter = any(k in (notes or "") for k in ["동일", "대상", "명단", "교직원", "교번", "연구자", "교수"])
    if not wants_reference_filter:
        return []

    from openpyxl import load_workbook
    for file_index, path in enumerate(request_file_paths, 1):
        if not is_spreadsheet_file(path):
            continue
        wb = load_workbook(path, data_only=True, read_only=True)
        for ws in wb.worksheets:
            for r in range(1, min(ws.max_row or 0, 10) + 1):
                for c in range(1, (ws.max_column or 0) + 1):
                    header = ws.cell(r, c).value
                    if not header:
                        continue
                    header_text = str(header).replace("\n", " ").strip()
                    if any(k in header_text for k in ["교직원", "교번", "직원번호", "사번"]):
                        wb.close()
                        return [{
                            "source_file_index": file_index,
                            "sheet": ws.title,
                            "column": c,
                            "data_start_row": r + 1,
                            "db_indices": [30, 34],
                            "match_type": "contains",
                            "label": header_text,
                        }]
                    if any(k in header_text for k in ["연구자", "교수", "성명", "이름", "발명자"]):
                        wb.close()
                        return [{
                            "source_file_index": file_index,
                            "sheet": ws.title,
                            "column": c,
                            "data_start_row": r + 1,
                            "db_indices": [29, 33],
                            "match_type": "contains",
                            "label": header_text,
                        }]
        wb.close()
    return []


def apply_reference_filters(rows, request_file_paths, reference_filters):
    if not request_file_paths or not reference_filters:
        return rows
    filtered = rows
    for reference_filter in reference_filters:
        source_index = int(reference_filter.get("source_file_index") or 1)
        if source_index < 1 or source_index > len(request_file_paths):
            continue
        values = set(_extract_reference_values(request_file_paths[source_index - 1], reference_filter))
        db_indices = []
        for db_idx in reference_filter.get("db_indices") or []:
            try:
                db_indices.append(int(db_idx))
            except Exception:
                pass
        if not values or not db_indices:
            continue

        before = len(filtered)
        match_type = reference_filter.get("match_type", "contains")

        def matches(row):
            for db_idx in db_indices:
                if db_idx >= len(row):
                    continue
                cell_value = _norm_match_value(row[db_idx])
                if not cell_value:
                    continue
                if match_type == "exact" and cell_value in values:
                    return True
                if match_type != "exact" and any(v in cell_value or cell_value in v for v in values):
                    return True
            return False

        filtered = [row for row in filtered if matches(row)]
        print(
            f"  요청파일 대상 필터 '{reference_filter.get('label', '')}' "
            f"→ {len(values)}개 값 / {before}행 → {len(filtered)}행"
        )
    return filtered


def format_cell_value(value, db_index: int):
    """셀 값 표시용 포맷."""
    if value is None:
        return ""
    # 날짜 컬럼
    if db_index in (1, 46, 62, 73, 84, 100, 101, 102):
        return to_yyyymmdd(value)
    # 숫자 컬럼 (금액 등)
    if db_index in (50, 52, 53, 60, 70, 71, 76, 77, 78, 79,
                    81, 82, 83, 85, 86, 87, 88, 89, 90, 91, 92, 93, 94):
        try:
            n = int(float(str(value)))
            return n
        except:
            pass
    return value


def run(master_path: str, year: int, output_path: str,
        notes: str, api_key: str, request_file_paths=None, request_file_names=None) -> dict:
    """
    메인 실행 함수.
    반환: {"count": N, "columns": [...], "total_income": N, "ai_notes": "..."}
    """
    print(f"[extract_columns] 마스터 DB 로딩: {master_path}")
    all_rows = load_master_db(master_path)
    request_file_paths = request_file_paths or []
    request_file_names = request_file_names or []
    print(f"  전체 행수: {len(all_rows)}")
    diagnostics = [{"step": "마스터 DB 로딩", "before": None, "after": len(all_rows)}]
    request_files_context = build_request_files_context(request_file_paths, request_file_names)
    if request_files_context.strip():
        print(f"  요청파일 참고 내용 분석: {len(request_files_context)}자")

    # ── 마스터 DB 유효성 검증 ────────────────────────────────────────
    if all_rows and len(all_rows[0]) < 50:
        raise ValueError(
            f"❌ 마스터 DB 파일이 올바르지 않습니다.\n"
            f"업로드한 파일의 열 수: {len(all_rows[0])}개 (정상: 105개)\n\n"
            f"👉 '기술이전총정리_날짜.xlsx' 파일을 마스터 DB 칸에 올려주세요.\n"
            f"   추출 결과 파일이나 양식 파일을 마스터 DB로 올리면 안 됩니다."
        )

    # ── Regex로 날짜범위·연구자 먼저 파싱 (AI 보조) ─────────────────
    import re as _re
    def _parse_notes(text):
        """notes에서 날짜범위·필터기준·연구자 목록을 regex로 추출."""
        res = {"date_start": None, "date_end": None, "date_mode": "both", "researchers": None}
        m = _re.search(r'(20\d{6})\s*[~\-]\s*(20\d{6})', text)
        if m:
            res["date_start"] = m.group(1)
            res["date_end"]   = m.group(2)
        else:
            m = _re.search(
                r'(20\d{2})\s*[.\-/년]\s*(\d{1,2})\s*[.\-/월]\s*(\d{1,2})\s*(?:[.일])?\s*[~\-–]\s*'
                r'(20\d{2})\s*[.\-/년]\s*(\d{1,2})\s*[.\-/월]\s*(\d{1,2})\s*(?:[.일])?',
                text,
            )
            if m:
                y1, mo1, d1, y2, mo2, d2 = m.groups()
                res["date_start"] = f"{int(y1):04d}{int(mo1):02d}{int(d1):02d}"
                res["date_end"] = f"{int(y2):04d}{int(mo2):02d}{int(d2):02d}"
        if "계약일" in text:
            res["date_mode"] = "contract"
        elif "입금일" in text:
            res["date_mode"] = "payment"
        rm = _re.search(r'(?:대상\s*)?연구자\s*[:：]\s*(.+)', text, _re.DOTALL)
        if rm:
            names = [n.strip() for n in _re.split(r'[,，/\n]', rm.group(1)) if n.strip()]
            if names:
                res["researchers"] = names
        return res

    regex_parsed = _parse_notes(notes or "")
    print(f"  Regex 파싱 - 날짜: {regex_parsed['date_start']}~{regex_parsed['date_end']}, "
          f"기준: {regex_parsed['date_mode']}, 연구자: {regex_parsed['researchers']}")

    # ── AI에게 열 및 필터 결정 요청 ────────────────────────────────
    print("[extract_columns] AI 분석 중...")
    try:
        ai_result = ask_openai(notes, year, api_key, request_files_context=request_files_context)
        columns           = ai_result.get("columns") or []
        filter_mode       = ai_result.get("filter_mode", "both")
        filter_year       = ai_result.get("filter_year") or year
        date_range        = ai_result.get("date_range")
        researcher_filter = ai_result.get("researcher_filter")
        reference_filters = ai_result.get("reference_filters") or []
        db_filters        = ai_result.get("db_filters") or []
        ai_notes          = ai_result.get("notes", "")
        print(f"  AI 결정 - 열: {columns}, 필터: {filter_mode}, 연도: {filter_year}")
        print(f"  날짜범위: {date_range}, 연구자: {researcher_filter}, DB조건: {len(db_filters)}개")
        print(f"  AI 설명: {ai_notes}")
    except Exception as e:
        print(f"  AI 분석 실패 ({e}), 기본값 사용")
        columns           = []
        filter_mode       = "both"
        filter_year       = year
        date_range        = None
        researcher_filter = None
        reference_filters = []
        db_filters        = []
        ai_notes          = "AI 분석 실패 → 기본 열 사용"

    # Regex 결과로 AI 누락분 보완
    if not date_range and regex_parsed["date_start"]:
        date_range = {"start": regex_parsed["date_start"], "end": regex_parsed["date_end"]}
    if not researcher_filter and regex_parsed["researchers"]:
        researcher_filter = regex_parsed["researchers"]
    if regex_parsed["date_mode"] != "both":
        filter_mode = regex_parsed["date_mode"]

    inferred_columns = infer_requested_columns(notes or "")
    if not columns and inferred_columns:
        columns = inferred_columns
        print(f"  요청사항 키워드 기반 열 보완 → {columns}")

    # 추출 항목 미지정이면 전체 열(0~104) 사용
    if not columns:
        columns = list(range(105))
        print("  추출 항목 미지정 → 전체 열(0~104) 추출")
    else:
        deduped = []
        for c in columns:
            if isinstance(c, int) and 0 <= c <= 104 and c not in deduped:
                deduped.append(c)
        columns = deduped

    # ── 날짜 필터링 ─────────────────────────────────────────────────
    if date_range and date_range.get("start") and date_range.get("end"):
        start_d = date_range["start"]
        end_d   = date_range["end"]
        before = len(all_rows)
        rows = filter_by_date_range(all_rows, start_d, end_d,
                                    contract_col=1, payment_col=73, mode=filter_mode)
        print(f"  날짜범위 필터 {start_d}~{end_d} ({filter_mode}) → {len(rows)}행")
        diagnostics.append({
            "step": f"날짜범위 필터 {start_d}~{end_d} ({filter_mode})",
            "before": before,
            "after": len(rows),
        })
    elif filter_mode == "none":
        rows = all_rows
        print(f"  필터 없음 → {len(rows)}행")
        diagnostics.append({"step": "필터 없음", "before": len(all_rows), "after": len(rows)})
    else:
        before = len(all_rows)
        rows = filter_by_year(all_rows, filter_year,
                              contract_col=1, payment_col=73, mode=filter_mode)
        print(f"  {filter_year}년 필터({filter_mode}) → {len(rows)}행")
        diagnostics.append({
            "step": f"{filter_year}년 필터({filter_mode})",
            "before": before,
            "after": len(rows),
        })

    # ── 연구자 필터링 ────────────────────────────────────────────────
    if researcher_filter:
        before = len(rows)
        rows = filter_by_researchers(rows, researcher_filter,
                                     inventor_col=29, co_inventor_col=33)
        print(f"  연구자 필터 {researcher_filter} → {before}행 → {len(rows)}행")
        diagnostics.append({
            "step": f"연구자 필터 {researcher_filter}",
            "before": before,
            "after": len(rows),
        })

    if not reference_filters:
        reference_filters = infer_reference_filters(request_file_paths, notes)
    if reference_filters:
        before = len(rows)
        rows = apply_reference_filters(rows, request_file_paths, reference_filters)
        diagnostics.append({
            "step": "요청파일 대상자/참고값 필터",
            "before": before,
            "after": len(rows),
        })

    if db_filters:
        before = len(rows)
        rows = apply_db_filters(rows, db_filters)
        diagnostics.append({
            "step": f"추가 DB 조건 필터 {len(db_filters)}개",
            "before": before,
            "after": len(rows),
        })

    # ── Excel 출력 ──────────────────────────────────────────────────
    wb = Workbook()
    ws = wb.active
    ws.title = f"{filter_year}년 추출"

    # 헤더 스타일
    header_fill  = PatternFill("solid", fgColor="1F3864")
    header_font  = Font(bold=True, color="FFFFFF", size=10)
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align   = Alignment(horizontal="left",   vertical="center")
    thin_border  = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"),  bottom=Side(style="thin"),
    )

    # 헤더 행
    for col_pos, db_idx in enumerate(columns, start=1):
        cell = ws.cell(row=1, column=col_pos, value=HEADER_MAP.get(db_idx, f"col_{db_idx}"))
        cell.fill      = header_fill
        cell.font      = header_font
        cell.alignment = center_align
        cell.border    = thin_border

    ws.row_dimensions[1].height = 30

    # 데이터 행
    for row_pos, row in enumerate(rows, start=2):
        for col_pos, db_idx in enumerate(columns, start=1):
            raw_val = row[db_idx] if db_idx < len(row) else None
            val = format_cell_value(raw_val, db_idx)
            cell = ws.cell(row=row_pos, column=col_pos, value=val)
            cell.border    = thin_border
            cell.alignment = left_align
            cell.font      = Font(size=9)

    # 열 너비 자동 조정 (최대 40)
    for col_pos, db_idx in enumerate(columns, start=1):
        header_len = len(HEADER_MAP.get(db_idx, "")) + 2
        max_len = header_len
        for row_pos in range(2, min(len(rows) + 2, 52)):  # 최대 50행 샘플
            cell_val = ws.cell(row=row_pos, column=col_pos).value
            if cell_val:
                max_len = max(max_len, min(len(str(cell_val)), 40))
        ws.column_dimensions[ws.cell(row=1, column=col_pos).column_letter].width = max_len + 2

    # 틀 고정
    ws.freeze_panes = "A2"

    # 총 입금액 계산 (76번 열이 포함된 경우)
    total_income = 0
    if 76 in columns:
        for row in rows:
            if 76 < len(row):
                total_income += safe_int(row[76])

    wb.save(output_path)
    print(f"[extract_columns] 저장 완료: {output_path} ({len(rows)}건)")
    zero_reasons = []
    if not rows:
        for item in diagnostics:
            before = item.get("before")
            after = item.get("after")
            if before and after == 0:
                zero_reasons.append(
                    f"{item.get('step')} 단계에서 {before}건이 0건으로 줄었습니다."
                )
        if not zero_reasons:
            zero_reasons.append("마스터 DB 또는 적용 조건에서 일치하는 행을 찾지 못했습니다.")

    return {
        "count": len(rows),
        "columns": columns,
        "column_validation": describe_columns(columns),
        "total_income": total_income,
        "filter_mode": filter_mode,
        "filter_year": filter_year,
        "ai_notes": ai_notes,
        "understanding": {
            "mode": "서식 없음 / 요청 항목 직접 추출",
            "request_files": request_file_names,
            "columns": describe_columns(columns),
            "date_range": date_range,
            "filter_mode": filter_mode,
            "filter_year": filter_year,
            "researcher_filter": researcher_filter,
            "reference_filters": reference_filters,
            "db_filters": db_filters,
        },
        "diagnostics": diagnostics,
        "zero_result_reasons": zero_reasons,
    }
